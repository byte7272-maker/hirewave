"""Interview prep engine + API + text extraction."""

from __future__ import annotations

from fastapi.testclient import TestClient

from jobsearch.api.app import create_app
from jobsearch.api.state import AppState
from jobsearch.engines.integration import MockTokenExchanger
from jobsearch.engines.interview import InterviewEngine
from jobsearch.models import Resume, ResumeSource
from jobsearch.textextract import extract_text


# --- engine -----------------------------------------------------------------
def test_generate_covers_categories(profile, matching_job):
    prep = InterviewEngine().generate(profile, job=matching_job, count=8)
    cats = {q.category.value for q in prep.questions}
    assert "intro" in cats and "closing" in cats
    assert "technical" in cats  # from job requirements / skills
    assert all(q.suggested_answer for q in prep.questions)  # every Q has an answer
    assert all(q.question for q in prep.questions)
    assert prep.job_posting_id == matching_job.id


def test_count_is_bounded(profile):
    assert len(InterviewEngine().generate(profile, count=100).questions) <= 12
    assert len(InterviewEngine().generate(profile, count=1).questions) >= 3


def test_grounded_on_resume_text(profile, matching_job):
    resume = Resume(
        user_id=profile.user_id,
        source=ResumeSource.UPLOADED,
        rendered_text="Sam Dev — built a payments microservice on FastAPI handling 2k req/s.",
    )
    prep = InterviewEngine().generate(profile, resume=resume, job=matching_job)
    assert prep.based_on_document is True
    assert prep.resume_id == resume.id


def test_no_document_uses_profile(profile):
    prep = InterviewEngine().generate(profile)
    assert prep.based_on_document is False


def test_gap_question_when_requirement_missing(profile, matching_job):
    matching_job.requirements = ["Python", "Rust"]  # Rust not in profile skills
    prep = InterviewEngine().generate(profile, job=matching_job, count=12)
    gap_qs = [q for q in prep.questions if q.category.value == "gap"]
    assert gap_qs and "Rust" in gap_qs[0].question


# --- text extraction --------------------------------------------------------
def test_extract_text_plain():
    assert "Python" in extract_text(b"# CV\nPython, FastAPI", filename="cv.md")


def test_extract_docx_roundtrip():
    import docx  # from the `documents` extra
    import io

    d = docx.Document()
    d.add_paragraph("Alex Rivera")
    d.add_paragraph("Senior Backend Engineer — Python, FastAPI, AWS")
    buf = io.BytesIO()
    d.save(buf)
    text = extract_text(buf.getvalue(), filename="resume.docx")
    assert "Alex Rivera" in text and "FastAPI" in text


def test_extract_unparseable_is_empty():
    assert extract_text(b"\x00\x01\x02not text", filename="x.bin") == ""


# --- API --------------------------------------------------------------------
def test_api_prep_grounded_on_uploaded_resume():
    client = TestClient(create_app(state=AppState(exchanger=MockTokenExchanger())))
    client.post(
        "/api/v1/auth/register",
        json={"email": "sam@demo.com", "password": "supersecret", "full_name": "Sam"},
    )
    tok = client.post(
        "/api/v1/auth/login", json={"email": "sam@demo.com", "password": "supersecret"}
    ).json()
    h = {"Authorization": f"Bearer {tok['access_token']}"}
    client.put("/api/v1/users/me", headers=h, json={"skills": ["Python", "FastAPI"]})

    # Upload a text résumé, then generate prep grounded on it.
    resume = client.post(
        "/api/v1/resumes/upload",
        headers=h,
        files={"file": ("cv.md", b"# Sam\nBuilt FastAPI services at Acme.", "text/markdown")},
    ).json()

    r = client.post(
        "/api/v1/interview/prep", headers=h, json={"resume_id": resume["id"], "count": 5}
    )
    assert r.status_code == 201
    prep = r.json()
    assert prep["based_on_document"] is True
    assert len(prep["questions"]) == 5
    assert all(q["suggested_answer"] for q in prep["questions"])


def test_api_prep_unknown_resume_404():
    client = TestClient(create_app(state=AppState(exchanger=MockTokenExchanger())))
    client.post(
        "/api/v1/auth/register",
        json={"email": "a@b.com", "password": "supersecret"},
    )
    tok = client.post(
        "/api/v1/auth/login", json={"email": "a@b.com", "password": "supersecret"}
    ).json()
    h = {"Authorization": f"Bearer {tok['access_token']}"}
    r = client.post("/api/v1/interview/prep", headers=h, json={"resume_id": "nope"})
    assert r.status_code == 404
